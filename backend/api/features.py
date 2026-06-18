"""
VerdictFlow — Standout Feature Endpoints

Bundles the hackathon "wow" features that build on the core pipeline:
  - POST /api/cases/{id}/chat            → "Ask this contract" RAG copilot
  - POST /api/cases/{id}/debate          → Agent boardroom cross-examination
  - GET  /api/cases/{id}/counter-draft   → Redlined .docx counter-draft (no LLM)
  - POST /api/cases/{id}/negotiation-email → Drafts the email to the counterparty

Every LLM-backed endpoint degrades gracefully (returns a sensible fallback)
so the demo never crashes on a transient model error or quota limit.
"""

import io
import logging

from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import Response
from pydantic import BaseModel

from agents.orchestrator import get_case, get_contract_text
from core.llm import llm_generate, parse_json_response

logger = logging.getLogger("verdictflow.features")

router = APIRouter()


# ── Models ───────────────────────────────────────────────────────────────────


class ChatRequest(BaseModel):
    question: str


# ── 1. Ask-this-Contract RAG copilot ─────────────────────────────────────────


@router.post("/api/cases/{case_id}/chat")
async def chat_with_contract(case_id: str, req: ChatRequest, request: Request):
    """Answer a natural-language question grounded in the contract via RAG."""
    packet = get_case(case_id)
    if not packet:
        raise HTTPException(404, f"Case '{case_id}' not found")

    vs = getattr(request.app.state, "vectorstore", None)
    doc_id = packet.contract.doc_id if packet.contract else None

    # Retrieve the most relevant chunks for grounding.
    passages: list[str] = []
    if vs and doc_id:
        try:
            hits = vs.search(query=req.question, doc_id=doc_id, top_k=5)
            passages = [h.get("text", "") for h in hits if h.get("text")]
        except Exception as e:
            logger.warning(f"Chat RAG search failed: {e}")

    # Fall back to the raw contract text if retrieval found nothing.
    context = "\n\n---\n\n".join(passages)
    if not context.strip():
        context = (get_contract_text(case_id) or "")[:8000]

    if not context.strip():
        return {"answer": "I don't have the contract text indexed yet. Please wait for analysis to finish.", "citations": []}

    system_prompt = (
        "You are a contract analysis assistant. Answer the user's question using ONLY the "
        "provided contract excerpts. Be specific and quote the relevant language. If the answer "
        "is not in the excerpts, say so plainly. Keep it concise (2-5 sentences)."
    )
    user_prompt = f"CONTRACT EXCERPTS:\n{context[:9000]}\n\nQUESTION: {req.question}\n\nAnswer:"

    try:
        answer = await llm_generate(system_prompt=system_prompt, user_prompt=user_prompt, max_tokens=1024, temperature=0.2)
        answer = (answer or "").strip() or "I couldn't generate an answer for that. Try rephrasing."
    except Exception as e:
        logger.warning(f"Chat LLM failed: {e}")
        answer = "The assistant is temporarily unavailable (model rate limit). Please try again shortly."

    return {"answer": answer, "citations": [p[:240] for p in passages[:3]]}


# ── 2. Agent Boardroom — cross-examination ───────────────────────────────────


@router.post("/api/cases/{case_id}/debate")
async def agent_debate(case_id: str):
    """
    Generate a multi-round cross-examination where the Red Team challenges the
    Clause Analyst's findings and the Adjudicator rules — showcasing genuine
    multi-agent reasoning. One LLM call produces the full structured transcript.
    """
    packet = get_case(case_id)
    if not packet:
        raise HTTPException(404, f"Case '{case_id}' not found")

    clause = [f"[{f.risk_level}] {f.category}: {f.explanation}" for f in packet.clause_findings[:6]]
    attacks = [f"[{a.severity}] {a.attack_vector}: {a.exploit_scenario}" for a in packet.red_team_attacks[:6]]
    compliance = [f"{c.regulation}: {c.finding}" for c in packet.compliance_checks[:4]]

    findings_block = (
        "CLAUSE ANALYST FINDINGS:\n" + ("\n".join(clause) or "none") +
        "\n\nRED TEAM ATTACKS:\n" + ("\n".join(attacks) or "none") +
        "\n\nCOMPLIANCE ISSUES:\n" + ("\n".join(compliance) or "none")
    )

    system_prompt = """You are simulating a contract-review boardroom where specialist AI agents
cross-examine each other's findings to reach a rigorous consensus. The agents are:
- Clause Analyst (defends its risk findings)
- Red Team (challenges and probes for missed exploits or overstated risk)
- Compliance (raises regulatory angles)
- Adjudicator (chairs, weighs the arguments, and rules)

Produce a punchy, realistic 5-7 message debate where agents actually CHALLENGE and respond to
each other (not just take turns). End with the Adjudicator's ruling.

Return ONLY a JSON array, each item:
{"speaker": "Clause Analyst|Red Team|Compliance|Adjudicator", "message": "1-3 sentences", "stance": "challenge|defend|concede|rule"}"""

    user_prompt = f"Here are the findings to debate:\n\n{findings_block}\n\nGenerate the boardroom cross-examination."

    rounds = []
    try:
        raw = await llm_generate(system_prompt=system_prompt, user_prompt=user_prompt, max_tokens=2048, temperature=0.6)
        parsed = parse_json_response(raw)
        if isinstance(parsed, list):
            for item in parsed:
                if isinstance(item, dict) and item.get("message"):
                    rounds.append({
                        "speaker": item.get("speaker", "Adjudicator"),
                        "message": str(item.get("message", "")),
                        "stance": item.get("stance", "defend"),
                    })
    except Exception as e:
        logger.warning(f"Debate LLM failed: {e}")

    if not rounds:
        # Deterministic fallback so the panel always shows something.
        rounds = [
            {"speaker": "Clause Analyst", "message": f"I flagged {len(packet.clause_findings)} risky clauses, with the most severe around liability and termination.", "stance": "defend"},
            {"speaker": "Red Team", "message": f"I found {len(packet.red_team_attacks)} exploitable angles those findings don't fully cover — the ambiguity is worse than rated.", "stance": "challenge"},
            {"speaker": "Compliance", "message": f"There are also {len(packet.compliance_checks)} regulatory exposures that compound the contractual risk.", "stance": "challenge"},
            {"speaker": "Adjudicator", "message": "Weighing all three: the risks are material and reinforce each other. The recommendation stands.", "stance": "rule"},
        ]

    return {"case_id": case_id, "rounds": rounds}


# ── 3. Counter-draft .docx (no LLM — reuses redline edits) ────────────────────


@router.get("/api/cases/{case_id}/counter-draft")
async def counter_draft(case_id: str):
    """Build a redlined .docx counter-draft from the Redline agent's edits."""
    packet = get_case(case_id)
    if not packet:
        raise HTTPException(404, f"Case '{case_id}' not found")

    from docx import Document
    from docx.shared import RGBColor, Pt

    doc = Document()
    title = doc.add_heading("VerdictFlow — Negotiation Counter-Draft", level=0)
    fname = packet.contract.filename if packet.contract else case_id
    doc.add_paragraph(f"Proposed redlines for: {fname}").italic = True
    doc.add_paragraph(
        f"{len(packet.redline_edits)} suggested edits. Struck-through red text is the original; "
        "green underlined text is the proposed replacement."
    )
    doc.add_paragraph("")

    if not packet.redline_edits:
        doc.add_paragraph("No redline edits were generated for this contract.")

    for i, edit in enumerate(packet.redline_edits, 1):
        h = doc.add_heading(f"Edit {i} · {edit.priority.value if hasattr(edit.priority, 'value') else edit.priority}", level=2)

        p_orig = doc.add_paragraph()
        r1 = p_orig.add_run("Original:  ")
        r1.bold = True
        r_old = p_orig.add_run(edit.original_text or "(missing clause)")
        r_old.font.strike = True
        r_old.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        p_new = doc.add_paragraph()
        r2 = p_new.add_run("Proposed:  ")
        r2.bold = True
        r_new = p_new.add_run(edit.suggested_text or "")
        r_new.font.underline = True
        r_new.font.color.rgb = RGBColor(0x1E, 0x8E, 0x3E)

        p_rat = doc.add_paragraph()
        rr = p_rat.add_run("Rationale: ")
        rr.bold = True
        rr.font.size = Pt(9)
        p_rat.add_run(edit.rationale or "").font.size = Pt(9)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    filename = f"verdictflow_counter_draft_{case_id[:8]}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── 4. Negotiation email drafter ─────────────────────────────────────────────


@router.post("/api/cases/{case_id}/negotiation-email")
async def negotiation_email(case_id: str):
    """Draft a professional email to the counterparty requesting the key changes."""
    packet = get_case(case_id)
    if not packet:
        raise HTTPException(404, f"Case '{case_id}' not found")

    top_edits = [f"- [{(e.priority.value if hasattr(e.priority, 'value') else e.priority)}] {e.rationale}" for e in packet.redline_edits[:6]]
    top_clauses = [f"- [{f.risk_level}] {f.category}: {f.explanation}" for f in packet.clause_findings[:5]]
    parties = packet.contract.parties if packet.contract else []
    counterparty = parties[1] if len(parties) > 1 else "the counterparty"

    system_prompt = (
        "You are a sharp but professional contracts lead. Draft a concise, courteous email to the "
        "counterparty requesting the most important contract changes. Be specific and firm, group the "
        "asks, and propose next steps. 150-220 words. Return ONLY the email text (with a Subject line)."
    )
    user_prompt = (
        f"Recipient: {counterparty}\n"
        f"Contract: {packet.contract.doc_type if packet.contract else 'agreement'}\n\n"
        f"Required edits:\n" + ("\n".join(top_edits) or "various") +
        f"\n\nKey risks:\n" + ("\n".join(top_clauses) or "various") +
        "\n\nDraft the email."
    )

    try:
        email = await llm_generate(system_prompt=system_prompt, user_prompt=user_prompt, max_tokens=900, temperature=0.4)
        email = (email or "").strip()
    except Exception as e:
        logger.warning(f"Negotiation email LLM failed: {e}")
        email = ""

    if not email:
        email = (
            f"Subject: Proposed revisions to our {packet.contract.doc_type if packet.contract else 'agreement'}\n\n"
            f"Hi {counterparty},\n\nThank you for sharing the draft. Before we can proceed, we'll need a few "
            "revisions to bring the terms in line with market standard — chiefly around liability, termination, "
            "and the renewal/pricing terms. I've attached a redlined counter-draft with the specifics.\n\n"
            "Could we set up 30 minutes this week to walk through them? I'm confident we can find balanced "
            "language quickly.\n\nBest regards,\n[Your name]"
        )

    return {"case_id": case_id, "email": email}
