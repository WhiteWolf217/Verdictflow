"""
VerdictFlow — Audit Packet DOCX Exporter

Renders a completed AuditPacket into a professional Word (.docx) audit report:
executive verdict, financial summary, clause flags, compliance issues, and
proposed redlines (original in red strikethrough, proposed in green), with a
tamper-evident SHA-256 footer. Returns the document as bytes (no disk I/O).
"""

import logging
from datetime import datetime, timezone
from io import BytesIO
from typing import Optional, TYPE_CHECKING

from docx import Document
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from models.schemas import AuditPacket

logger = logging.getLogger("verdictflow.audit_packager")

RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)

_REC_LABEL = {
    "approve": "APPROVE",
    "approve_with_changes": "APPROVE WITH CHANGES",
    "do_not_sign": "DO NOT SIGN",
}


def _trunc(text: str, n: int = 300) -> str:
    text = (text or "").strip()
    return text if len(text) <= n else text[: n - 1] + "…"


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def build_audit_docx(packet: "AuditPacket", latest_hash: Optional[str] = None) -> bytes:
    """Render an AuditPacket to a .docx report and return the bytes."""
    doc = Document()

    filename = packet.contract.filename if packet.contract else "Unknown contract"
    doc_type = packet.contract.doc_type if packet.contract else "unknown"

    # ── Title ──
    doc.add_heading("VerdictFlow Contract Audit Report", level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        f"{filename}  ·  type: {doc_type}  ·  "
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    sub_run.italic = True
    sub_run.font.color.rgb = GREY
    doc.add_paragraph(f"Case ID: {packet.case_id}")

    # ── 1. Executive Verdict ──
    doc.add_heading("1. Executive Verdict", level=1)
    rec = packet.verdict_recommendation.value if packet.verdict_recommendation else None
    rec_para = doc.add_paragraph()
    rec_run = rec_para.add_run(
        f"Recommendation: {_REC_LABEL.get(rec, 'PENDING')}  "
        f"(confidence {packet.confidence_score:.0%})"
    )
    rec_run.font.bold = True
    rec_run.font.size = Pt(13)
    if rec == "do_not_sign":
        rec_run.font.color.rgb = RED
    elif rec == "approve":
        rec_run.font.color.rgb = GREEN
    doc.add_paragraph(packet.verdict or "No verdict was generated.")

    # ── 2. Financial Risk Summary ──
    doc.add_heading("2. Financial Risk Summary", level=1)
    summary = packet.risk_summary
    doc.add_paragraph(
        f"Total estimated exposure: ${summary.get('total_financial_exposure', 0):,.0f}"
    )
    if packet.financial_risks:
        _add_table(
            doc,
            ["Category", "Exposure", "Score", "Explanation"],
            [
                [
                    r.category,
                    f"${r.exposure_amount:,.0f} {r.currency}" if r.exposure_amount else "—",
                    f"{r.risk_score:.2f}",
                    _trunc(r.explanation, 200),
                ]
                for r in packet.financial_risks
            ],
        )
    else:
        doc.add_paragraph("No financial risks were identified.")

    # ── 3. Clause Flags ──
    doc.add_heading("3. Clause Flags", level=1)
    if packet.clause_findings:
        _add_table(
            doc,
            ["Category", "Risk", "Explanation"],
            [
                [f.category, f.risk_level.value.upper(), _trunc(f.explanation, 280)]
                for f in packet.clause_findings
            ],
        )
    else:
        doc.add_paragraph("No clause-level risks were flagged.")

    # ── 4. Compliance Issues ──
    doc.add_heading("4. Compliance Issues", level=1)
    if packet.compliance_checks:
        _add_table(
            doc,
            ["Regulation", "Status", "Finding", "Remediation"],
            [
                [
                    c.regulation,
                    c.status.value.replace("_", " ").title(),
                    _trunc(c.finding, 200),
                    _trunc(c.remediation, 200),
                ]
                for c in packet.compliance_checks
            ],
        )
    else:
        doc.add_paragraph("No compliance issues were identified.")

    # ── 5. Proposed Redlines ──
    doc.add_heading("5. Proposed Redlines", level=1)
    if packet.redline_edits:
        for i, edit in enumerate(packet.redline_edits, 1):
            head = doc.add_paragraph()
            head_run = head.add_run(
                f"Redline {i} — {edit.priority.value.upper()}"
            )
            head_run.font.bold = True

            para = doc.add_paragraph()
            orig = para.add_run("ORIGINAL: " + _trunc(edit.original_text, 400))
            orig.font.color.rgb = RED
            orig.font.strike = True
            para.add_run("\n")
            prop = para.add_run("PROPOSED: " + _trunc(edit.suggested_text, 400))
            prop.font.color.rgb = GREEN

            rationale = doc.add_paragraph()
            r_run = rationale.add_run("Rationale: " + _trunc(edit.rationale, 300))
            r_run.italic = True
            r_run.font.color.rgb = GREY
    else:
        doc.add_paragraph("No redline edits were proposed.")

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    f_run = footer.add_run(
        f"Tamper-evident SHA-256: {latest_hash or 'n/a'}\n"
        f"Generated by VerdictFlow · {datetime.now(timezone.utc).isoformat()}"
    )
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = GREY

    buf = BytesIO()
    doc.save(buf)
    logger.info(f"📄 Built audit DOCX for case '{packet.case_id}' ({buf.tell()} bytes)")
    return buf.getvalue()
