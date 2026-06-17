"""
VerdictFlow — Contract Document Parser

Supports PDF (via PyMuPDF) and DOCX (via python-docx) extraction.
Returns structured ParsedDocument with page-level text and metadata.
"""

import logging
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from pydantic import BaseModel, Field

logger = logging.getLogger("verdictflow.parser")


# ── Data Models ──────────────────────────────────────────────────────────────


class PageContent(BaseModel):
    """Content extracted from a single page/section."""

    page_number: int
    text: str
    headings: list[str] = Field(default_factory=list)
    char_count: int = 0


class ParsedDocument(BaseModel):
    """Complete parsed document with metadata."""

    doc_id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_type: str  # "pdf" or "docx"
    page_count: int
    total_chars: int
    pages: list[PageContent]
    raw_text: str  # Full concatenated text
    parsed_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


# ── PDF Parser ───────────────────────────────────────────────────────────────


def parse_pdf(file_path: str, filename: Optional[str] = None) -> ParsedDocument:
    """
    Parse a PDF file using PyMuPDF (fitz).

    Extracts text page-by-page, identifies headings by font size heuristic,
    and returns a structured ParsedDocument.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    pages: list[PageContent] = []
    all_text_parts: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")

        # Extract headings by looking for larger font sizes
        headings: list[str] = []
        blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    # Heuristic: text with font size >= 14 is likely a heading
                    if span["size"] >= 14 and span["text"].strip():
                        headings.append(span["text"].strip())

        page_content = PageContent(
            page_number=page_num + 1,
            text=text,
            headings=headings,
            char_count=len(text),
        )
        pages.append(page_content)
        all_text_parts.append(text)

    raw_text = "\n\n".join(all_text_parts)

    result = ParsedDocument(
        filename=filename or Path(file_path).name,
        file_type="pdf",
        page_count=len(doc),
        total_chars=len(raw_text),
        pages=pages,
        raw_text=raw_text,
    )

    doc.close()
    logger.info(f"📄 Parsed PDF '{result.filename}': {result.page_count} pages, {result.total_chars} chars")
    return result


# ── DOCX Parser ──────────────────────────────────────────────────────────────


def parse_docx(file_path: str, filename: Optional[str] = None) -> ParsedDocument:
    """
    Parse a DOCX file using python-docx.

    Extracts paragraphs with style-aware metadata (headings by style name).
    Groups content into logical "pages" based on section breaks or heading groups.
    """
    from docx import Document

    doc = Document(file_path)
    paragraphs_text: list[str] = []
    headings: list[str] = []
    current_section_lines: list[str] = []
    pages: list[PageContent] = []
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ""

        # Detect headings
        if "heading" in style_name:
            # Start a new logical page on major headings
            if current_section_lines and ("heading 1" in style_name or "heading 2" in style_name):
                pages.append(PageContent(
                    page_number=page_num,
                    text="\n".join(current_section_lines),
                    headings=headings.copy(),
                    char_count=sum(len(l) for l in current_section_lines),
                ))
                page_num += 1
                current_section_lines = []
                headings = []

            headings.append(text)

        current_section_lines.append(text)
        paragraphs_text.append(text)

    # Add remaining content as final page
    if current_section_lines:
        pages.append(PageContent(
            page_number=page_num,
            text="\n".join(current_section_lines),
            headings=headings,
            char_count=sum(len(l) for l in current_section_lines),
        ))

    raw_text = "\n\n".join(paragraphs_text)

    result = ParsedDocument(
        filename=filename or Path(file_path).name,
        file_type="docx",
        page_count=len(pages),
        total_chars=len(raw_text),
        pages=pages,
        raw_text=raw_text,
    )

    logger.info(f"📄 Parsed DOCX '{result.filename}': {result.page_count} sections, {result.total_chars} chars")
    return result


# ── Plain Text Parser ────────────────────────────────────────────────────────


def parse_txt(file_path: str, filename: Optional[str] = None) -> ParsedDocument:
    """
    Parse a plain-text (.txt) contract.

    Reads the file as UTF-8 (ignoring undecodable bytes), splits into logical
    "pages" on form-feed characters if present (otherwise a single page), and
    extracts ALL-CAPS / numbered lines as headings.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        raw_text = f.read()

    # Split on form-feeds into pages; fall back to one page for the whole doc.
    page_texts = [p for p in raw_text.split("\f")] if "\f" in raw_text else [raw_text]

    heading_re = re.compile(r"^(?:\d+\.[\d.]*\s+[A-Z].{2,60}|[A-Z][A-Z \-/]{4,60})\s*$")
    pages: list[PageContent] = []
    for i, text in enumerate(page_texts):
        headings = [
            line.strip()
            for line in text.splitlines()
            if heading_re.match(line.strip())
        ]
        pages.append(PageContent(
            page_number=i + 1,
            text=text,
            headings=headings,
            char_count=len(text),
        ))

    result = ParsedDocument(
        filename=filename or Path(file_path).name,
        file_type="txt",
        page_count=len(pages),
        total_chars=len(raw_text),
        pages=pages,
        raw_text=raw_text,
    )
    logger.info(
        f"📄 Parsed TXT '{result.filename}': {result.page_count} page(s), {result.total_chars} chars"
    )
    return result


# ── Dispatcher ───────────────────────────────────────────────────────────────


# Extensions VerdictFlow can parse. Kept here so the API layer can validate
# uploads against a single source of truth.
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt")


def parse_contract(file_path: str, filename: Optional[str] = None) -> ParsedDocument:
    """
    Parse a contract file (PDF, DOCX, or plain text).

    Automatically detects file type from extension and dispatches
    to the appropriate parser. Raises ValueError on unsupported types and
    wraps parser-level failures with file context.
    """
    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".pdf":
            return parse_pdf(file_path, filename)
        elif ext in (".docx", ".doc"):
            return parse_docx(file_path, filename)
        elif ext == ".txt":
            return parse_txt(file_path, filename)
        else:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
            )
    except ValueError:
        raise
    except Exception as e:  # noqa: BLE001 — surface a clean, contextual error
        raise ValueError(
            f"Failed to parse '{filename or file_path}' ({ext}): {e}"
        ) from e
